import re
from typing import List, BinaryIO
from docx import Document

# Use absolute imports
from src.DTOs.models import (
    CongaMergeField,
    CongaControlTag,
    TextSegment,
    ParsedTemplateElement
)

class DocxParserError(Exception):
    """Custom exception for DOCX parsing errors."""
    pass

# Regex to find Conga tags: {{...}}
# It captures the content inside the double curly braces.
CONGA_TAG_REGEX = r"\{\{([^}]+)\}\}" # For python: r"{{([^}]+)}}"

def _parse_tag_content(full_tag: str, inner_content: str) -> ParsedTemplateElement:
    """
    Parses the inner content of a Conga tag to determine if it's a merge field or control tag.
    """
    inner_content_stripped = inner_content.strip()
    if ":" in inner_content_stripped:
        parts = inner_content_stripped.split(":", 1)
        control_type = parts[0].strip()
        parameter = parts[1].strip() if len(parts) > 1 else None
        return CongaControlTag(
            original_tag=full_tag,
            control_type=control_type,
            parameter=parameter
        )
    else:
        return CongaMergeField(
            original_tag=full_tag,
            field_name=inner_content_stripped
        )

def _extract_elements_from_text(text: str) -> List[ParsedTemplateElement]:
    """
    Extracts all Conga tags and text segments from a given string of text.
    """
    elements: List[ParsedTemplateElement] = []
    last_end = 0
    
    # Find all tag matches in the text
    for match in re.finditer(CONGA_TAG_REGEX, text):
        # Add any text before the current tag as a TextSegment
        if match.start() > last_end:
            text_before = text[last_end:match.start()]
            if text_before.strip():  # Only add non-whitespace text segments
                elements.append(TextSegment(
                    original_tag=text_before,
                    content=text_before
                ))
        
        # Process the tag itself
        full_tag = match.group(0)      # e.g., {{FieldName}} or {{TableStart:Contacts}}
        inner_content = match.group(1)  # e.g., FieldName or TableStart:Contacts
        
        try:
            element = _parse_tag_content(full_tag, inner_content)
            elements.append(element)
        except Exception as e: 
            print(f"Warning: Could not parse tag content for '{full_tag}'. Error: {e}")
            # If parsing fails, keep the original tag as a text segment
            elements.append(TextSegment(
                original_tag=full_tag,
                content=full_tag
            ))
        
        last_end = match.end()
    
    # Add any remaining text after the last tag
    if last_end < len(text):
        remaining_text = text[last_end:]
        if remaining_text.strip():  # Only add non-whitespace text segments
            elements.append(TextSegment(
                original_tag=remaining_text,
                content=remaining_text
            ))
    
    return elements

def extract_elements_from_docx(file_path: str) -> List[ParsedTemplateElement]:
    """
    Parses a .docx Conga template to extract merge fields and control tags.

    Args:
        file_path: The absolute path to the .docx template file.

    Returns:
        A list of ParsedTemplateElement objects (CongaMergeField or CongaControlTag).
        Returns an empty list if no tags are found or if the document is empty.
        
    Raises:
        DocxParserError: If the file cannot be read or if there's an issue with the docx format.
    """
    try:
        document = Document(file_path)
    except FileNotFoundError:
        raise DocxParserError(f"Error: DOCX template file not found at {file_path}")
    except Exception as e: # Catches other errors from python-docx like bad format
        raise DocxParserError(f"Error: Could not read or open DOCX template file {file_path}. Details: {e}")

    all_elements: List[ParsedTemplateElement] = []

    # Process paragraphs
    for para in document.paragraphs:
        if para.text.strip():  # Only process non-empty paragraphs
            all_elements.extend(_extract_elements_from_text(para.text))
            # Add a newline to separate paragraphs (optional, adjust as needed)
            all_elements.append(TextSegment(
                original_tag="\n",
                content="\n"
            ))

    # Process tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                # Each cell can contain multiple paragraphs
                for para_in_cell in cell.paragraphs:
                    if para_in_cell.text.strip():
                        all_elements.extend(_extract_elements_from_text(para_in_cell.text))
                # Add a tab or other separator between cells (optional)
                all_elements.append(TextSegment(
                    original_tag="\t",
                    content="\t"
                ))
            # Add a newline after each row
            all_elements.append(TextSegment(
                original_tag="\n",
                content="\n"
            ))
    
    return all_elements

if __name__ == '__main__':
    import os
    PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    DUMMY_DOCX_FILE = os.path.join(PROJECT_ROOT, "input_data", "conga_template.docx")

    if not os.path.exists(DUMMY_DOCX_FILE) or os.path.getsize(DUMMY_DOCX_FILE) == 0:
        try:
            print(f"Creating a simple dummy DOCX for testing at {DUMMY_DOCX_FILE}...")
            doc = Document()
            # Note: f-string curly braces need to be doubled to be literal {{}}
            doc.add_paragraph(f"Hello {{{{Name}}}}, this is a test.")
            doc.add_paragraph("{{IF:ShowDetails}}Details here.{{ENDIF}}")
            table_obj = doc.add_table(rows=1, cols=2) # Renamed variable
            table_obj.cell(0,0).text = "{{TableStart:Contacts}}"
            table_obj.cell(0,1).text = "{{ContactName}}"
            doc.add_paragraph("{{TableEnd:Contacts}}")
            doc.save(DUMMY_DOCX_FILE)
            print("Dummy DOCX created.")
        except Exception as e:
            print(f"Could not create dummy DOCX: {e}")

    print(f"Attempting to parse: {DUMMY_DOCX_FILE}")
    if os.path.exists(DUMMY_DOCX_FILE):
        try:
            elements = extract_elements_from_docx(DUMMY_DOCX_FILE)
            if elements:
                print(f"\nSuccessfully extracted {len(elements)} elements from DOCX:")
                for i, elem in enumerate(elements):
                    print(f"  Element {i+1}: {elem.model_dump_json()}") # Using model_dump_json for Pydantic v2
                    if isinstance(elem, CongaMergeField):
                        print(f"    Type: MergeField, Name: {elem.field_name}")
                    elif isinstance(elem, CongaControlTag):
                        print(f"    Type: ControlTag, Control: {elem.control_type}, Param: {elem.parameter}")
            else:
                print("No elements extracted or parsing returned None/empty list.")
        except DocxParserError as e:
            print(f"DOCX Parsing Test Failed: {e}")
        except Exception as e:
            print(f"An unexpected error occurred during DOCX parsing testing: {e}")
    else:
        print(f"Test DOCX file {DUMMY_DOCX_FILE} not found and could not be created. Skipping DOCX parser test.")

    print("\nAttempting to parse non_existent_template.docx:")
    try:
        extract_elements_from_docx("non_existent_template.docx")
    except DocxParserError as e:
        print(f"Successfully caught expected error: {e}")
